import os
from pathlib import Path
import logging
from typing import Dict, List
import re
import fitz 
from docx import Document as DocxDocument
from openpyxl import load_workbook
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    pdfplumber = None

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Handle document processing for Chinese and English documents.
    Supporting type: 
    DOCX (with tables and images): .docx,
    PDF (with tables and images): .pdf,
    Excel: .xlsx, .xls,
    text files: .txt, 
    markdown files: .md,
    Images (with OCR): .png, .jpg, .jpeg, .bmp, .gif, .tiff, .webp
    """

    def __init__(self, upload_file_dir: str = "processed_texts", enable_ocr: bool = False):
        """
        Initialize document processor.
        Args: 
            upload_file_dir (str): Directory containing documents to process.
            enable_ocr (bool): Whether to enable OCR for image files.
        """
        self.upload_file_dir = Path(upload_file_dir)
        self.upload_file_dir.mkdir(exist_ok=True)
        
        if enable_ocr:
            try:
                from rapidocr_onnxruntime import RapidOCR
                self.ocr_engine = RapidOCR()
            except ImportError:
                logger.warning("RapidOCR not available, OCR features disabled")
                self.ocr_engine = None
        else:
            self.ocr_engine = None

    def get_pdf_page_count(self, file_path: str) -> int:
        """Get the number of pages in a PDF file."""
        try:
            with fitz.open(file_path) as doc:
                return len(doc)
        except Exception as e:
            logger.error(f"Error getting PDF page count: {e}")
            return 0

    def convert_pdf_to_images(self, file_path: str, max_pages: int = 5) -> List:
        """
        Convert PDF pages to images for multimodal processing.
        Args:
            file_path: Path to PDF
            max_pages: Maximum number of pages to convert
        Returns:
            List of PIL Images
        """
        from PIL import Image
        import io
        
        images = []
        try:
            with fitz.open(file_path) as doc:
                num_pages = min(len(doc), max_pages)
                for i in range(num_pages):
                    page = doc[i]
                    # Higher DPI (2.0) for better text recognition by the VL model
                    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                    img_data = pix.tobytes("png")
                    images.append(Image.open(io.BytesIO(img_data)))
            return images
        except Exception as e:
            logger.error(f"Error converting PDF to images: {e}")
            return []

    def process_documents(self) -> Dict[str, str]:
        """
        Process all documents in the specified directory and clean extracted text.
        Returns: 
            Dict[str, str]: A dictionary mapping file paths to cleaned text content.
        """
        processed_texts = {}
        
        # Walk through the directory
        for root, dirs, files in os.walk(self.upload_file_dir):
            for file in files:
                file_path = os.path.join(root, file)
                file_extension = os.path.splitext(file)[1].lower()
                
                try:
                    content = ""
                    if file_extension == '.docx':
                        logger.info(f"Processing DOCX: {file_path}")
                        content = self.extract_text_from_docx_to_markdown(file_path)
                    
                    elif file_extension == '.pdf':
                        logger.info(f"Processing PDF: {file_path}")
                        content = self.extract_text_from_pdf_to_markdown(file_path)
                    
                    elif file_extension in ['.xlsx', '.xls']:
                        logger.info(f"Processing Excel: {file_path}")
                        content = self.extract_text_from_excel_to_markdown(file_path)
                    
                    elif file_extension in ['.txt', '.md']:
                        logger.info(f"Processing Text/MD: {file_path}")
                        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                            content = f.read()

                    elif file_extension in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp']:
                        logger.info(f"Processing Image: {file_path}")
                        content = self.extract_information_from_image(file_path)
                    
                    else:
                        logger.warning(f"Skipping unsupported file: {file_path}")
                        continue

                    # Clean the extracted text
                    processed_texts[file_path] = self.clean_markdown_text(content)

                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    processed_texts[file_path] = f"[Processing failed: {os.path.basename(file_path)} - {str(e)}]"

        return processed_texts
    
    def extract_information_from_image(self, image_path: str) -> str:
        """
        Extract text from image using OCR.
        Args: 
            image_path (str): Path to the image file.
        Returns:
            str: Extracted text from the image.
        """
        if not self.ocr_engine:
            return "[OCR engine not initialized]"
        try:
            result, _ = self.ocr_engine(image_path)
            if result:
                ocr_text = "\n".join([line[1] for line in result])
                return "[No text detected in image]"
        except Exception as e:
            logger.error(f"image extracted failed for {image_path}: {e}")
            return f"[Image Extraction failed: {e}]"

    def clean_markdown_text(self, text: str) -> str:
        """
        Clean text while preserving Markdown structure.
        Args:
            text (str): Raw extracted text.
        Returns:
            str: Cleaned text.
        """
        if not text or not text.strip():
            return ""
        
        try:
            # Remove control characters (except newlines/tabs)
            text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
            
            # Normalize whitespace
            # Replace multiple spaces with single space, but keep indentation for lists could be tricky
            # For safety, we just collapse runs of spaces that aren't newlines
            text = re.sub(r'[ \t]+', ' ', text)
            
            # Fix multiple newlines (max 2)
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error cleaning text: {e}")
            return text

    def extract_text_from_docx_to_markdown(self, file_path: str) -> str:
        """
        Extract text from DOCX to Markdown, including OCR for images.
        Args:
            file_path (str): Path to the DOCX file.
        Returns:
            str: Extracted Markdown text.
        """
        try:
            doc = DocxDocument(file_path)
            lines = []
            filename = os.path.splitext(os.path.basename(file_path))[0]
            lines.append(f"# {filename}\n")

            # Helper to extract and OCR images from an XML element
            def extract_images_from_element(element):
                if not self.ocr_engine:
                    return []
                
                # XML Namespaces for docx
                ns = {
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                }
                
                extracted = []
                # Find all blip elements (images)
                # Note: This finds images in the element's scope (paragraph or cell)
                blips = element.findall('.//a:blip', ns)
                
                for blip in blips:
                    rId = blip.get(f"{{{ns['r']}}}embed")
                    if rId and rId in doc.part.related_parts:
                        try:
                            part = doc.part.related_parts[rId]
                            # Run OCR on the image blob
                            result, _ = self.ocr_engine(part.blob)
                            if result:
                                ocr_text = " ".join([line[1] for line in result])
                                if ocr_text.strip():
                                    extracted.append(f"> [Image OCR]: {ocr_text}")
                        except Exception as e:
                            logger.warning(f"Error OCRing image in DOCX: {e}")
                return extracted
            
            for para in doc.paragraphs:
                text = para.text.strip()
                
                # Check for images in this paragraph
                image_texts = extract_images_from_element(para._element)
                
                if not text and not image_texts:
                    continue
                
                # Simple header detection based on style or content
                style_name = para.style.name if para.style else None
                if style_name and 'Heading' in style_name:
                    level = 1
                    try:
                        # Extract level from "Heading 1", "Heading 2" etc.
                        level = int(style_name.split()[-1])
                    except:
                        level = 2
                    lines.append(f"{'#' * level} {text}")
                elif self._is_likely_header(text):
                    lines.append(f"## {text}")
                else:
                    lines.append(text)
                
                # Append any image text found in the paragraph
                if image_texts:
                    lines.append("\n" + "\n".join(image_texts) + "\n")
            
            # Handle tables in DOCX
            for table in doc.tables:
                lines.append("\n") # Spacing before table
                rows = []
                for row in table.rows:
                    # Extract text AND images from cells
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ')
                        
                        # Check for images in the cell
                        cell_images = extract_images_from_element(cell._element)
                        if cell_images:
                            # Append image text to cell text
                            img_text_joined = " ".join(cell_images)
                            cell_text += f" {img_text_joined}"
                            
                        row_cells.append(cell_text)
                    rows.append(row_cells)
                
                if rows:
                    # Construct Markdown table
                    header = rows[0]
                    lines.append("| " + " | ".join(header) + " |")
                    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                    for row in rows[1:]:
                        lines.append("| " + " | ".join(row) + " |")
                lines.append("\n")

            return '\n'.join(lines)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return f"[Error extracting DOCX: {e}]"

    def extract_text_from_pdf_to_markdown(self, file_path: str) -> str:
        """
        Extract text from PDF using pdfplumber (table) with enhanced image OCR support
        Args:
            file_path (str): Path to the PDF file.
        Returns:
            str: Extracted Markdown text.
        """
        try:
            if PDFPLUMBER_AVAILABLE:
                return self._extract_pdf_with_plumber(file_path)
            else:
                return self._extract_pdf_with_fitz(file_path)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return f"[Error extracting PDF: {e}]"

    def _extract_pdf_with_plumber(self, file_path: str) -> str:
        """
        PDF extraction using pdfplumber (Text + Tables) with enhanced Image OCR support.
        Strategy: Iterate pages, extract tables, extract images (OCR), extract text, and attempt to linearize them based on vertical position.
        Args: 
            file_path (str): Path to the PDF file.
        Returns: 
            str: Extracted Markdown text.
        """
        markdown_output = []
        filename = os.path.splitext(os.path.basename(file_path))[0]
        markdown_output.append(f"# {filename}\n")

        with pdfplumber.open(file_path) as pdf, fitz.open(file_path) as fitz_doc:
            for page_idx, page in enumerate(pdf.pages):
                markdown_output.append(f"## Page {page_idx + 1}\n")
                
                # Get the corresponding PyMuPDF page for accurate image rendering
                fitz_page = fitz_doc[page_idx]

                # 1. Extract Tables
                # pdfplumber extracts tables as list of list of strings
                tables = page.find_tables()
                table_bboxes = [t.bbox for t in tables]
                
                extracted_items = [] # List of (top_y, type, content)

                # Process Tables
                for table in tables:
                    # Convert table to markdown
                    data = table.extract()
                    if not data:
                        continue
                    
                    # Clean data (None -> "")
                    data = [[cell if cell is not None else "" for cell in row] for row in data]
                    # Filter empty rows
                    data = [row for row in data if any(c.strip() for c in row)]
                    
                    if data:
                        md_table = []
                        # Header
                        headers = [str(c).replace('\n', ' ') for c in data[0]]
                        md_table.append("| " + " | ".join(headers) + " |")
                        md_table.append("| " + " | ".join(["---"] * len(headers)) + " |")
                        # Rows
                        for row in data[1:]:
                            clean_row = [str(c).replace('\n', ' ') for c in row]
                            md_table.append("| " + " | ".join(clean_row) + " |")
                        
                        extracted_items.append((table.bbox[1], 'table', "\n".join(md_table)))

                # 2. Extract Images & OCR
                if self.ocr_engine:
                    page_area = page.width * page.height
                    text_char_count = len(page.chars)

                    for img in page.images:
                        # pdfplumber image objects have x0, top, x1, bottom
                        x0, top, x1, bottom = img['x0'], img['top'], img['x1'], img['bottom']
                        
                        # Filter out very small images (icons, decorations)
                        if (bottom - top) < 20 or (x1 - x0) < 20:
                            continue
                        
                        # Optimization: Skip OCR for full-page background images if text is already present
                        img_area = (x1 - x0) * (bottom - top)
                        if page_area > 0 and (img_area / page_area) > 0.8 and text_char_count > 50:
                            continue

                        # Check if image is substantially inside a table
                        mid_x = (x0 + x1) / 2
                        mid_y = (top + bottom) / 2
                        is_in_table = False
                        for (tx0, ttop, tx1, tbottom) in table_bboxes:
                            if tx0 <= mid_x <= tx1 and ttop <= mid_y <= tbottom:
                                is_in_table = True
                                break
                        if is_in_table:
                            continue
                            
                        try:
                            # Use fitz to render the specific area
                            # pdfplumber coords are typically PDF points (72 DPI based). Fitz uses same system by default.
                            rect = fitz.Rect(x0, top, x1, bottom)
                            # Get pixmap of the rect
                            pix = fitz_page.get_pixmap(clip=rect)
                            img_data = pix.tobytes("png")
                            
                            # Run OCR
                            result, _ = self.ocr_engine(img_data)
                            if result:
                                ocr_texts = [line[1] for line in result]
                                full_ocr_text = "\n".join(ocr_texts)
                                if full_ocr_text.strip():
                                    formatted_text = f"\n> [Image OCR]: {full_ocr_text.replace('\n', ' ')}\n"
                                    extracted_items.append((top, 'image', formatted_text))
                        except Exception as e:
                            logger.warning(f"Image OCR failed on page {page_idx}: {e}")

                # 3. Extract Text (Filtering out table areas to avoid duplication)
                # We define a filtering function for extracts_words
                def not_inside_table(obj):
                    # obj has x0, top, x1, bottom
                    obj_x0 = obj['x0']
                    obj_top = obj['top']
                    obj_x1 = obj['x1']
                    obj_bottom = obj['bottom']
                    
                    # Check if center of word is inside any table bbox
                    mid_x = (obj_x0 + obj_x1) / 2
                    mid_y = (obj_top + obj_bottom) / 2
                    
                    for (tx0, ttop, tx1, tbottom) in table_bboxes:
                        if tx0 <= mid_x <= tx1 and ttop <= mid_y <= tbottom:
                            return False
                    return True

                # Extract words that are NOT inside tables
                text_content = page.filter(not_inside_table).extract_text()
                
                if not tables:
                    # Simple case: Just text
                    if text_content:
                        extracted_items.append((0, 'text', text_content))
                else:
                    # Complex case: Interleave
                    # We can use .extract_words() to rebuild lines and paragraphs with Y coordinates
                    words = page.filter(not_inside_table).extract_words()
                    if words:
                        # Group words into lines
                        lines = {} # y_coord -> list of words
                        for word in words:
                            # Round top to group loosely aligned words (tolerance ~3px)
                            top_key = round(word['top'] / 3) * 3
                            if top_key not in lines:
                                lines[top_key] = []
                            lines[top_key].append(word)
                        
                        # Sort lines by Y
                        sorted_y = sorted(lines.keys())
                        for y in sorted_y:
                            line_words = sorted(lines[y], key=lambda w: w['x0'])
                            line_text = " ".join([w['text'] for w in line_words])
                            extracted_items.append((y * 3, 'text', line_text)) # Use approx Y

                # 4. Fallback Page OCR (Only if completely empty)
                if not extracted_items and self.ocr_engine:
                    # Render page to image for OCR
                    try:
                        # Use fitz to render image as pdfplumber image handling can be complex
                        pix = fitz_page.get_pixmap()
                        img_data = pix.tobytes("png")
                        # RapidOCR can take bytes
                        result, _ = self.ocr_engine(img_data)
                        if result:
                            ocr_text = "\n".join([line[1] for line in result])
                            extracted_items.append((0, 'ocr', ocr_text))
                    except Exception as e:
                        logger.warning(f"Page {page_idx} empty and OCR failed: {e}")

                # Sort all items by vertical position (top)
                extracted_items.sort(key=lambda x: x[0])
                
                # Combine
                for _, type_, content in extracted_items:
                    if type_ == 'table':
                        markdown_output.append("\n" + content + "\n")
                    else:
                        markdown_output.append(content)
                        
                markdown_output.append("\n") # Page break spacing

        return "\n".join(markdown_output)

    def _extract_pdf_with_fitz(self, file_path: str) -> str:
        """
        Fallback extraction using PyMuPDF.
        Args: 
            file_path (str): Path to the PDF file.
        Returns:
            str: Extracted Markdown text.
        """
        doc = fitz.open(file_path)
        markdown_lines = []
        filename = os.path.splitext(os.path.basename(file_path))[0]
        markdown_lines.append(f"# {filename}\n")
        
        for page_num, page in enumerate(doc):
            markdown_lines.append(f"## Page {page_num + 1}\n")
            
            # Extract text
            text = page.get_text()
            if text.strip():
                markdown_lines.append(text)
            
            # Extract images for OCR if needed
            if self.ocr_engine:
                image_list = page.get_images()
                for img in image_list:
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n - pix.alpha < 3: # CMYK or other
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        img_data = pix.tobytes()
                        
                        result, _ = self.ocr_engine(img_data)
                        if result:
                            ocr_text = "\n".join([line[1] for line in result])
                            markdown_lines.append(f"\n[Image Content]:\n{ocr_text}\n")
                    except Exception:
                        pass
        
        return "\n".join(markdown_lines)

    def extract_text_from_excel_to_markdown(self, file_path: str) -> str:
        """
        Extract text from Excel to Markdown.
        Args: 
            file_path (str): Path to the Excel file.
        Returns:
            str: Extracted Markdown text.
        """
        try:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            lines = []
            filename = os.path.splitext(os.path.basename(file_path))[0]
            lines.append(f"# {filename}\n")
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines.append(f"## Sheet: {sheet_name}\n")
                
                rows = []
                for row in ws.rows:
                    # Convert to string and handle None
                    row_data = [str(cell.value) if cell.value is not None else "" for cell in row]
                    # Keep row if it has any content
                    if any(cell.strip() for cell in row_data):
                        rows.append(row_data)
                
                if rows:
                    # Create markdown table
                    # Make sure all rows have same length
                    max_len = max(len(r) for r in rows)
                    rows = [r + [""] * (max_len - len(r)) for r in rows]
                    
                    header = rows[0]
                    lines.append("| " + " | ".join([h.replace('\n', ' ') for h in header]) + " |")
                    lines.append("| " + " | ".join(["---"] * max_len) + " |")
                    
                    for row in rows[1:]:
                        clean_row = [c.replace('\n', ' ') for c in row]
                        lines.append("| " + " | ".join(clean_row) + " |")
                lines.append("\n")
                
            return "\n".join(lines)
        except Exception as e:
            logger.error(f"Excel extraction error: {e}")
            return f"[Error extracting Excel: {e}]"

    def _is_likely_header(self, text: str) -> bool:
        """
        Heuristic to detect headers in plain text.
        Agrs: 
            text (str): Text line to evaluate.
        Returns:
            bool: True if likely a header, False otherwise.
        """
        if len(text) > 80: 
            return False
        
        # Patterns: 
        # 1. Numbered start: "1. Introduction", "2.1 Analysis"
        # 2. Chinese Section: "第一章", "一、"
        # 3. Short capitalized text (English)
        
        if re.match(r'^[\d\.]+\s', text):
            return True
        if re.match(r'^[IVX]+\.', text):
            return True
        if re.match(r'^[第][一二三四五六七八九十\d]+[章节]', text):
            return True
            
        return False


